import docx
import sys
sys.stdout.reconfigure(encoding='utf-8')

print("=" * 80)
print("DOCUMENT 1: Learnavo_Student_Module_Implementation.docx")
print("=" * 80)
doc1 = docx.Document("Learnavo_Student_Module_Implementation.docx")
for p in doc1.paragraphs:
    if p.text.strip():
        print(p.text)

print("\n\n")
print("=" * 80)
print("DOCUMENT 2: Learnavo_Student_Module_v2.docx")
print("=" * 80)
doc2 = docx.Document("Learnavo_Student_Module_v2.docx")
for p in doc2.paragraphs:
    if p.text.strip():
        print(p.text)
